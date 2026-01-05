import re
from docx import Document
from docx.shared import Pt

def add_formatted_text(paragraph, text):
    """
    Helper to parse Markdown bold/italic and add as runs to a paragraph.
    """
    # Regex to find **bold**, *italic*, or plain text
    # This splits the text into parts while keeping the delimiters
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold: remove ** and set bold=True
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic: remove * and set italic=True
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Plain text
            paragraph.add_run(part)

def create_admission_report(university, program, content, output_path):
    doc = Document()
    
    # Styled Title
    doc.add_heading(f'Admissions Strategy: {university}', 0)
    doc.add_heading(f'Program: {program}', level=1)
    
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 1. Handle Headings
        if line.startswith('###'):
            clean_text = line.replace('###', '').strip()
            doc.add_heading(clean_text, level=3)
        elif line.startswith('##'):
            clean_text = line.replace('##', '').strip()
            doc.add_heading(clean_text, level=2)
        elif line.startswith('#'):
            clean_text = line.replace('#', '').strip()
            doc.add_heading(clean_text, level=1)
            
        # 2. Handle Lists
        elif line.startswith(('* ', '+ ', '- ')):
            clean_line = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, clean_line)
            
        # 3. Handle Numbered Lists (e.g., 1. Research...)
        elif re.match(r'^\d+\.', line):
            clean_line = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, clean_line)
            
        # 4. Normal Paragraphs
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            
    doc.save(output_path)
    return output_path